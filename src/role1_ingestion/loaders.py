from pathlib import Path


def load_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def load_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def load_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    pieces = [p.text for p in doc.paragraphs]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                pieces.append(cell.text)

    return "\n".join(pieces)


def load_html(path: Path) -> str:
    from bs4 import BeautifulSoup

    html = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(html, "html.parser")

    for tag in soup.find_all(["script", "style"]):
        tag.decompose()

    return soup.get_text(separator="\n")


LOADERS = {
    ".txt": load_txt,
    ".md": load_txt,
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".html": load_html,
    ".htm": load_html,
}


def load_document(path: Path) -> str:
    """Read any supported file and return its plain text."""
    loader = LOADERS.get(path.suffix.lower())
    if loader is None:
        raise ValueError(f"Unsupported file type: {path.suffix}")
    return loader(path)